# -*- coding: utf-8 -*-
"""
Created on Tue Aug 21 09:40:03 2018

@author: Christos.TSELAS , Data Analyst - Machine Learning Engineer; 
ctselas@gmail.com
"""
import classification_performace_ct
from docx import Document 
from sklearn.model_selection import cross_val_predict, cross_val_score
import numpy as np
from sklearn import tree
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import scale
import graphviz
import operator as op
from docx.shared import Inches
from matplotlib import pyplot as plt
import seaborn as sns
import os
import random
random.seed(1)

class Classification(object):
    def __init__(self,train_x, train_y, name):
        'Initiliaze the parameters'
        self.names = train_x.columns
        self.train_x = train_x
        self.train_y = np.array(train_y)
        self.name = name
        
        #create a folder for Results if doesnt exist
        self.cwd = os.getcwd() 
        self.cwd = self.cwd + '\\Results'
        if not os.path.exists(self.cwd):
            os.makedirs(self.cwd)
        
        
    def decision_trees(self, max_depth = 10):
        'Train and Visualize a Decision Tree'
        # create predictions usins CV to explore performance
        clf = tree.DecisionTreeClassifier(max_depth = max_depth)
        pred_y = cross_val_predict(clf,self.train_x, self.train_y, cv = 10)
        # visualize and save performance
        meanauc = classification_performace_ct.Classification_Performance(self.train_y, pred_y, self.name, flag = 1).meanauc
    
        #train DT with all data and create the visualization
        clf.fit(self.train_x, self.train_y)
        dot_data = tree.export_graphviz(clf, out_file=None, 
                         impurity = False, 
                         feature_names=self.names,  
                         class_names=['Normal', 'Anomalies'],  
                         filled=True, rounded=True,  
                         special_characters=True) 
        graph = graphviz.Source(dot_data,format='png')
        graph
        #save in a pdf and in the report with the rest results
        graph.render(self.cwd + "\\Decision_Tree" + self.name[:-4]) 
#        self.document.add_heading('Decision Tree with max depht = 10 and mean AUC (10 fold CV) = ' + str(round(np.mean(validated),2))) 
        document = Document(self.cwd + '\\Performance of Dataset ' + self.name[:-4] + '.docx')
        document.add_picture(self.cwd + "\\Decision_Tree" + self.name[:-4] + '.png',width=Inches(7), height=Inches(4))
        document.save(self.cwd + '\\Performance of Dataset ' + self.name[:-4] + '.docx')
        return clf, meanauc
        

    def random_forests_all(self,  interval = [5, 35]):
        'Perform 3 different experiments using RF'
        n_trees, maxscores = self.random_forest_check_trees(interval) #check the best size (number of trees) of RF 
        importances, clf, meanauc = self.random_forest(n_trees) #find the most important features
        self.class_feature_importance(importances) #find the most important features for each class
        return clf, meanauc
        
        
    def random_forest_check_trees(self, interval = [5, 35]):
        'Train several RF using 10-fold CV and find the proper number of trees'
        scores = {}
        maxscores = 0
        indexmaxscores = 0
        for val in range(interval[0], interval[1]):
            clf = RandomForestClassifier(n_estimators = val)
            validated = cross_val_score(clf, self.train_x, self.train_y, cv = 10, scoring = 'roc_auc')
            scores[val] = list(validated)
            if np.mean(scores[val]) > maxscores:
                maxscores = np.mean(scores[val])
                indexmaxscores = val

        #results visualization
        sorted_keys, sorted_vals = zip(*sorted(scores.items(), key=op.itemgetter(1)))#sort keys and values together
        plt.figure()
        #almost verbatim from question
        plt.title('Performance of Random Forests in dataset: ' + self.name)
        sns.set(context='notebook', style='whitegrid')
        sns.boxplot(data=sorted_vals, width=.18)
        plt.xlabel('Number of Trees')
        plt.ylabel('AUC for 10-fold Cross Validation')
        plt.ylim([0,1.1])
        #category labels
        plt.xticks(plt.xticks()[0], sorted_keys)
        
        result ='Using 10-fold Cross Validation we compared RF with number of trees in' \
                                  + str(interval) + '\nBest performance:\nMean AUC = ' \
                                  + str(maxscores) + '\nNumber of trees = ' \
                                  + str(indexmaxscores)
        print (result)
        #results saving 
        plt.savefig(self.cwd + '\\RF_Performance_' + self.name[:-4] + '.png', bbox_inches='tight')
#        self.document.add_picture(self.cwd + '\\RF_Performance_' + self.name[:-4] + '.png',width=Inches(7), height=Inches(4))
#        self.document.save(self.cwd + '\\Report_' + self.name[:-4] + '.docx')
        return indexmaxscores, maxscores
    
    
    def random_forest(self, n_trees = 20):
        'Train a RF with all data and reveal the most important features'
        clf = RandomForestClassifier(n_estimators = n_trees)
        pred_y = cross_val_predict(clf,self.train_x, self.train_y, cv = 10)
        # visualize and save performance
        meanauc = classification_performace_ct.Classification_Performance(self.train_y, pred_y, self.name, flag = 1).meanauc
        
        clf = RandomForestClassifier(n_estimators = n_trees)
        clf.fit(self.train_x, self.train_y)
        importances = clf.feature_importances_
    
        indices = np.argsort(importances)[::-1] #sort the features
        feature_list = self.train_x.columns[indices]  #names of features.
        ff = np.array(feature_list)
#        # Print the feature ranking
#        print("Feature ranking:")
#        self.document.add_heading('Feature ranking:')
#        paragraph = self.document.add_paragraph()#create paragraph to save this print  
#        for f in range(self.train_x.shape[1]):
#            print("%d. feature %d (%f) name: %s" % (f + 1, indices[f], importances[indices[f]], ff[indices[f]]))
#            paragraph.add_run("\n%d. feature %d (%f) name: %s" % (f + 1, indices[f], importances[indices[f]], ff[indices[f]]))
        
        # Plot the feature importances of the forest
        plt.figure()
        plt.rcParams['figure.figsize'] = [16, 6]
        plt.title('RF_Feature importances, dataset: ' + self.name + 'Number of Trees' + str(n_trees))
        plt.bar(range(self.train_x.shape[1]), importances[indices],
                color="r", align="center")
        plt.xticks(range(self.train_x.shape[1]), ff[indices], rotation=90)
        plt.xlim([-1, self.train_x.shape[1]])
        plt.gcf().subplots_adjust(bottom=0.15)
        plt.show();
        
        # save the results
        plt.savefig(self.cwd + '\\RF_Feature_importance' + self.name[:-4] + '.png', bbox_inches='tight')    
#        self.document.add_picture(self.cwd + '\\RF_Feature_importance' + self.name[:-4] + '.png',width=Inches(7), height=Inches(4))
#        self.document.save(self.cwd + '\\Report_' + self.name[:-4] + '.docx')
        return importances , clf, meanauc
    
    
    def class_feature_importance_help(self, feature_importances):
        'To get the importance according to each class:'
        N, M = self.train_x.shape
        X = scale(self.train_x)
        out = {}
        for c in set(self.train_y):
            out[c] = dict(
                    zip(range(N), np.mean(X[self.train_y == c, :], axis=0)*feature_importances))
        return out    
    

    def class_feature_importance(self, importances):
        'Plot the feature importances of the forest for each class'
        result = self.class_feature_importance_help(importances)        
        indices = np.argsort(importances)[::-1]
        ff = np.array(self.train_x.columns[indices])  #names of features.
        
        titles = ["Class 0", "Class 1"]
        fig = plt.figure(figsize=(10,30))
        fig.suptitle('Feature Importance per Class, dataset: ' + self.name)
        for t, i in zip(titles, range(len(result))):
            plt.subplot(2,1,i+1)
            plt.title(t)
            plt.bar(range(len(result[i])), result[i].values(),
                      color="r", align="center")
            if i == 1:    
                plt.xticks(range(len(result[i])), ff[list(result[i].keys())], rotation=90)
                plt.gcf().subplots_adjust(bottom=0.15)
            
            plt.xlim([-1, len(result[i])])
            plt.show();
                
        plt.savefig(self.cwd + '\\Feature_importance_per_class' + self.name[:-4] + '.png', bbox_inches='tight')
#        self.document.add_picture(self.cwd + '\\Feature_importance_per_class' + self.name[:-4] + '.png',width=Inches(7), height=Inches(4))
#        self.document.save(self.cwd + '\\Report_' + self.name[:-4] + '.docx')
