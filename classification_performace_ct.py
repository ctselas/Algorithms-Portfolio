# -*- coding: utf-8 -*-
"""
Created on Mon Sep 17 09:11:25 2018

@author: Christos.TSELAS , Data Analyst - Machine Learning Engineer; 
ctselas@gmail.com
"""
from sklearn.metrics import confusion_matrix
from matplotlib import pyplot as plt
import seaborn as sns
from sklearn import metrics
from docx import Document
import os


class Classification_Performance(object):
    def __init__(self,target, predicted, name, flag = 0):
        'Initiliaze the parameters'
        self.target = target
        self.predicted = predicted
        self.name = name
        #create a folder for Results if doesnt exist
        self.cwd = os.getcwd() 
        self.cwd = self.cwd + '\\Results'
        if not os.path.exists(self.cwd):
            os.makedirs(self.cwd)
        #create a word document for reporting the results
        self.document = Document()
        self.document.save(self.cwd + '\\Performance of Dataset ' \
                               + self.name[:-4] + '.docx')
        self.document.add_heading('Results for Decision Tree with\
                                  max depht 7 and 10 fold CV') 
        
        self.paragraph = self.document.add_paragraph()#create paragraph to save this print  
        self.paragraph.text = '\n=====================================\n'
        self.paragraph.text = self.paragraph.text + 'Performance of \
                                Dataset:' + self.name +'\n'
        print('\n=====================================')
        print ('Performance of Dataset: ', self.name)
        
        if flag == 1: 
            self.meanauc = self.area_under_the_ROC()
            self.confusionMatrix()
            
        self.document.save(self.cwd + '\\Performance of Dataset ' + self.name[:-4] + '.docx')
        
        
    def confusionMatrix(self):
         'Visualization of Confusion Matrix'
         #Computation 
         conf_matrix = confusion_matrix(self.target, self.predicted)
         #Plot the confusion matrix 
         plt.figure(figsize=(8, 6))
         cmap=plt.cm.Blues
         sns.heatmap(conf_matrix,annot=True,fmt="d" , cmap=cmap)
         plt.title("Confusion matrix")
         plt.ylabel('True class')
         plt.xlabel('Predicted class')
         plt.show();
         # save the results
         plt.savefig(self.cwd + '\\Confusion_Matrix_' + self.name[:-4] + '.png', bbox_inches='tight')    
         self.document.add_picture(self.cwd + '\\Confusion_Matrix_' + self.name[:-4] + '.png')


    def area_under_the_ROC(self):
        'Generates FPR, TPR, AUC and visualize ROC'
        fpr, tpr, thresholds = metrics.roc_curve(self.target, self.predicted)
        roc_auc = metrics.auc(fpr, tpr)
        print ('False Positive Rate (FPR):',  round(fpr[1],2))
        print ('True Positive Rate (TPR):',  round(tpr[1],2))
        print ('Mean Area Under the ROC:', round(roc_auc,2))
        self.paragraph.text = self.paragraph.text + 'False Positive Rate' \
                '(FPR = FP/N = ‘Normals predicted as anomalies’/ ‘Normals’):' + '\n'
        self.paragraph.text = self.paragraph.text + '-> FPR = ' \
                + str(round(fpr[1],2)) + '\n'
        self.paragraph.text = self.paragraph.text + 'True Positive Rate' \
                '(TPR= TP/P = ‘Anomalies predicted correctly’ / ‘Anomalies’):'  + '\n'
        self.paragraph.text = self.paragraph.text + '-> TPR = '\
                + str(round(tpr[1],2)) + '\n'
        self.paragraph.text = self.paragraph.text + 'Mean Area Under the ROC: '\
                + str(round(roc_auc,2)) + '\n'
  
        fig, ax = plt.subplots()        
        plt.title('Receiving Operating Characteristic')
        ax.plot(fpr, tpr, 'b', label = 'ROC curve of class "anomaly" (AUC = %0.2f' % roc_auc + ')')
        textstr = '(' + str(round(fpr[1],2)) + ',' + str(round(tpr[1],2)) + ')'
        ax.text(fpr[1] , tpr[1] + 0.05 , textstr, transform=ax.transAxes, 
                fontsize=14,verticalalignment='top', color ='b')   
        plt.legend(loc = 'lower right')
        plt.scatter(fpr[1],tpr[1], c='b')
        plt.plot([0, 1], [0, 1], 'r--')
        plt.xlim([0, 1])
        plt.ylim([0, 1])
        plt.ylabel('True Positive Rate')
        plt.xlabel('False Positive Rate')
        plt.show();
        plt.savefig(self.cwd + '\\Receiving_Operating_Curve_' + self.name[:-4] + '.png', bbox_inches='tight')    
        self.document.add_picture(self.cwd + '\\Receiving_Operating_Curve_' + self.name[:-4] + '.png')
        return roc_auc


