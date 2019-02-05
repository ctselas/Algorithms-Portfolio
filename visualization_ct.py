# -*- coding: utf-8 -*-
"""
Created on Fri Aug 17 16:17:04 2018

@author: Christos.TSELAS , Data Analyst - Machine Learning Engineer; 
ctselas@gmail.com
"""

import numpy as np
from sklearn.decomposition import PCA
import matplotlib.pyplot as plt
import math
from matplotlib import cm as cm
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os
import tkinter as tk
#retrieve screen size (to be used for images) 
root = tk.Tk()
width = root.winfo_screenwidth()
height = root.winfo_screenheight()

class Visualization(object):
    def __init__(self,data, name, flag = 1):
        'Initiliaze the parameters'
        self.flag = flag
        self.names = data.columns
        self.name = name
        #create a folder for Results if doesnt exist
        self.cwd = os.getcwd() 
        self.cwd = self.cwd + '\\Results'
        if not os.path.exists(self.cwd):
            os.makedirs(self.cwd)
        
        if self.flag == 1 :
            self.features = np.matrix(data.iloc[:,:-1])
            self.target = np.matrix(data.iloc[:,-1]).T
        else: 
            self.features = np.matrix(data)
            
        loop = 'y'
        while loop == 'y':
           answer = input ('Hi! We have these 5 different options for visualization:\n\
                           1. Principal Component Analysis, and visualization of first 2 PCs\n \
                           2. Visualization of each feature and relation with target\n \
                           3. Plot of correlation matrix\n \
                           4. Plot the distribution of each feature\n \
                           5. Check for Outliers\n \
                           0. All the above(and some extra details) in a .doc file\n \
                           Choose one of the options pressing its number\n')
           if answer == '1':
               self.pca_vis()
           elif answer == '2':
               self.feature_vis()
           elif answer == '3':
               self.correlation_matrix()
           elif answer == '4':
               self.distributions()
           elif answer == '5':
               self.outliers_check()
           elif answer == '0':  
               self.all_functions()
           else:
               print ('Your input didnt match with any of the expected inputs \n')
           
           loop = input('Do you want a visualization ? Enter "y" for yes or "n" for no\n')
           while loop not in ('y','n'):
                   loop = input('Do you want a visualization ? Enter "y" for yes or "n" for no\n')
        print ('I hope the figures helped you :)')
                    
            
    def pca_vis(self):
        'Principal Component Analysis, and visualization of first 2 PCs'
        self.features = self.features-np.matlib.repmat(np.mean(self.features,axis=0),self.features.shape[0], 1)
        pca = PCA(n_components=2)
        pca.fit(self.features)
        reduced_features = pca.transform(self.features)
        fig = plt.figure(figsize=(width/100., height/100.), dpi=100)

        if self.flag == 1 :
            unique = list(set(np.ravel(self.target)))
            colors = [plt.cm.rainbow(float(i)/(max(unique)+1)) for i in unique]
            for i, u in enumerate(unique):
                xi = [reduced_features[j,0] for j  in range(len(reduced_features[:,0])) if np.ravel(self.target)[j] == u]
                yi = [reduced_features[j,1] for j  in range(len(reduced_features[:,0])) if np.ravel(self.target)[j] == u]
                plt.scatter(xi, yi, c=colors[i], label='Class ' + str(u))                
                        
            plt.xlabel('1st Principal Component')
            plt.ylabel('2nd Principal Component')
            plt.title('PCA Visualization')
            plt.legend()
            plt.show()
                
        else:
            plt.scatter(reduced_features[:,0], reduced_features[:,1])
            plt.xlabel('1st Principal Component')
            plt.ylabel('2nd Principal Component')
            plt.title('PCA Visualization')
            figManager = plt.get_current_fig_manager()
            figManager.window.showMaximized() 
            fig.tight_layout()
        output =  'Percentage of variance explained by each of the selected components:\n' + str(pca.explained_variance_ratio_)    
        print(output)      
        plt.savefig(self.cwd + '\\Principal_Component_Analysis_' + self.name[:-4] + '.png', bbox_inches='tight')    

        return output
                
                
    def feature_vis(self):
        'Visualization of each feature and relation with target'    
        number_of_features = self.features.shape[1]
        number_of_figugres = math.ceil(number_of_features/9.)
        for i in range(int(number_of_figugres)):
            fig = plt.figure(figsize=(width/100., height/100.), dpi=100)
            fig.suptitle('Feature Visualization')
            count = 0
            for j in range(i*9,min(9*(i+1),number_of_features),1):
                count  += 1 
                plt.subplot(3,3,count)
                plt.scatter(range(self.features.shape[0]),np.ravel(self.features[:,j]),marker='x',color='b',label='Feature Value')
                if self.flag == 1:
                    plt.scatter(range(self.target.shape[0]),np.ravel(self.target[:,-1]),marker='o',color='r',label='Class Value')

                plt.title(self.names[j])
                
            plt.legend(loc=1,bbox_to_anchor=(1,1),fontsize=12)
            fig.text(0.5, 0.04, 'Value', ha='center',fontsize=18)
            fig.text(0.04, 0.5, 'Sample', va='center', rotation='vertical',fontsize=18)
            figManager=plt.get_current_fig_manager()#Maximize the Window
            figManager.window.showMaximized() 
            fig.tight_layout()
            fig = plt.gcf()#Acquire figure before the Image Loads
            plt.show(block=False) #Show the Image (but don't block the flow)
            name = self.cwd +'\\Feature_Visualization_'+self.name[:-4]+'_'+str(i)+'.png'
            fig.savefig(name) #Save the image

        
    def correlation_matrix(self):
        'Plot of correlation matrix'
        if self.flag == 1 :
            df = np.concatenate((self.features,self.target[:,-1]),axis =1 ).T
            number_of_features = self.features.shape[1] + 1
        else:
            df = self.features.T
            number_of_features = self.features.shape[1]

        fig = plt.figure(figsize=(width/100., height/100.), dpi=100)
        fig.tight_layout()
        ax1 = fig.add_subplot(111)
        cmap = cm.get_cmap('seismic', 30)
        sns.heatmap(np.corrcoef(df), cmap=cmap, center=0,square=True, linewidths=.5, cbar_kws={"shrink": .5})
        ax1.grid(True)
        plt.title('Feature Correlation')
        plt.xticks(np.arange(number_of_features),self.names,fontsize=8, rotation=90)
        plt.yticks(np.arange(number_of_features),self.names,fontsize=8)
        figManager=plt.get_current_fig_manager()#Maximize the Window
        figManager.window.showMaximized() 
        plt.show()
        plt.savefig(self.cwd + '\\Correlation_Matrix_' + self.name[:-4] + '.png', bbox_inches='tight')    
    
    
    def distributions(self):
        'Plot the distribution of each feature'
        number_of_features = self.features.shape[1]
        number_of_figugres = math.ceil(number_of_features/9.)
        for i in range(int(number_of_figugres)):
            fig = plt.figure(figsize=(width/100., height/100.), dpi=100)
            fig.suptitle('Feature Distribution')
            count = 0
            for j in range(i*9,min(9*(i+1),number_of_features),1):
                count  += 1 
                plt.subplot(3,3,count)
                sns.distplot(self.features[:,j],kde=True)
                plt.title(self.names[j])
                
            fig.text(0.5, 0.04, 'Value', ha='center',fontsize=18)
            fig.text(0.04, 0.5, 'Sample', va='center', rotation='vertical',fontsize=18)
            figManager=plt.get_current_fig_manager()#Maximize the Window
            figManager.window.showMaximized() 
            plt.savefig(self.cwd +'\\Feature_Distribution_'+self.name[:-4]+'_'+str(i)+'.png', bbox_inches='tight')    
            
    def outliers_check(self):
        'Check if our data have outliers'
        number_of_features = self.features.shape[1]
        number_of_figugres = math.ceil(number_of_features/27.)
        for i in range(int(number_of_figugres)):
            fig = plt.figure(figsize=(width/100., height/100.), dpi=100)
            fig.suptitle('Check for Outliers')
            count = 0
            for j in range(i*27,min(27*(i+1),number_of_features),1):
                count  += 1 
                plt.subplot(3,9,count)
                sns.set_style('whitegrid')
                sns.boxplot(self.features[:,j],color='green',orient='v')
                plt.title(self.names[j])
                
            fig.text(0.5, 0.04, 'Value', ha='center',fontsize=18)
            fig.text(0.04, 0.5, 'Sample', va='center', rotation='vertical',fontsize=18)
            figManager=plt.get_current_fig_manager()#Maximize the Window
            figManager.window.showMaximized() 
            plt.savefig(self.cwd +'\\Check_for_Outliers_'+self.name[:-4]+'_'+str(i)+'.png', bbox_inches='tight')    
            
    
    def all_functions(self):
           'Call all functions and create a report'
           #create a word document for reporting the results
           document = Document()
           document.add_heading('Initial investigations on Data') 
           
           #PCA
           output = self.pca_vis()
           paragraph = document.add_paragraph()#create paragraph to save this print  
           runner = paragraph.add_run('\n=====================================\n')
           runner.bold = True
           runner = paragraph.add_run('Principal Component Analysis\n')
           runner.bold = True
           document.add_picture(self.cwd + '\\Principal_Component_Analysis_' + self.name[:-4] + '.png', width=Inches(6.5))
           paragraph.add_run(output)
           
           #Feature Visualization
           self.feature_vis()
           paragraph = document.add_paragraph()#create paragraph to save this print  
           runner = paragraph.add_run('\n=====================================\n')
           runner.bold = True
           runner = paragraph.add_run('Features Visualization\n')
           runner.bold = True
           number_of_features = self.features.shape[1]
           number_of_figugres = math.ceil(number_of_features/9.)
           for i in range(int(number_of_figugres)):
                   document.add_picture(self.cwd +'\\Feature_Visualization_'+self.name[:-4]+'_'+str(i)+'.png', width=Inches(6.5))
                   
           #Correlation Matrix
           self.correlation_matrix()
           paragraph = document.add_paragraph()#create paragraph to save this print  
           runner = paragraph.add_run('\n=====================================\n')
           runner.bold = True
           runner = paragraph.add_run('Correlation Matrix\n')
           runner.bold = True
           paragraph.add_run('Explanation:\nRed shades represents positive linear correlation while blue shades represents linear negative correlation.')
           document.add_picture(self.cwd + '\\Correlation_Matrix_' + self.name[:-4] + '.png', width=Inches(6.5))
          
           #Feature Distribution
           self.distributions()
           paragraph = document.add_paragraph()#create paragraph to save this print  
           runner = paragraph.add_run('\n=====================================\n')
           runner.bold = True
           runner = paragraph.add_run('Features Distribution\n')
           runner.bold = True
           paragraph.add_run("Explanation:\nA box plot (or box-and-whisker plot) shows the distribution of quantitative data in a way that facilitates comparisons between variables.The box shows the quartiles of the dataset while the whiskers extend to show the rest of the distribution. The box plot (a.k.a. box and whisker diagram) is a standardized way of displaying the distribution of data based on the five number summary:\n*Minimum\n*First quartile\n*Median \n*Third quartile \n*Maximum.\nIn the simplest box plot the central rectangle spans the first quartile to the third quartile (the interquartile range or IQR). A segment inside the rectangle shows the median and “whiskers” above and below the box show the locations of the minimum and maximum.")
           runner = paragraph.add_run('\nOutliers are either 3×IQR or more above the third quartile or 3×IQR or more below the first quartile.')
           runner.bold = True
           number_of_features = self.features.shape[1]
           number_of_figugres = math.ceil(number_of_features/9.)
           for i in range(int(number_of_figugres)):
                   document.add_picture(self.cwd +'\\Feature_Distribution_'+self.name[:-4]+'_'+str(i)+'.png', width=Inches(6.5))
                   
           #Check Outliers
           self.outliers_check()
           paragraph = document.add_paragraph()#create paragraph to save this print  
           runner = paragraph.add_run('\n=====================================\n')
           runner.bold = True
           runner = paragraph.add_run('Check for Outliers\n')
           runner.bold = True
           number_of_features = self.features.shape[1]
           number_of_figugres = math.ceil(number_of_features/27.)
           for i in range(int(number_of_figugres)):
                   document.add_picture(self.cwd +'\\Check_for_Outliers_'+self.name[:-4]+'_'+str(i)+'.png', width=Inches(6.5))
                   
           document.save(self.cwd + '\\Exploratory Data Analysis of ' + self.name[:-4] + '.docx')

    
    
    
    
    
    
if __name__ == "__main__" :  
    
    # initial parameters
    directory = r'\\aeroconseil.com\Shares\FPC\FDY\09-DATA_SCIENCES\DataAnalytics-Internal\Dionisis.K\Type3 Anomalies Data with labels'        
    name = 'p1813V0035_labeled.csv'
#    sample_size = 800
    # create the Preproces class
    d1 =  preprocessing_ct.Preprocess(name, directory)            
    d1.read_data()
    
    d1.delete_first_column()
    d1.delete_constant_rows()
    d1.delete_constant_columns()
    data = d1.train  
    m1 = Visualization(data[0:1000],name,flag = 0)
#    m1.pca_vis()
#    m1.feature_vis()
#    m1.correlation_matrix()
#    